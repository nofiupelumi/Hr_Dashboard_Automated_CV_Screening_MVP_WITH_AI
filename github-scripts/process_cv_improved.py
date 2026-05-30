#!/usr/bin/env python3
"""
Improved CV Processing Script for GitHub Actions
Handles authentication and better error handling for Laravel integration
"""

import os
import sys
import requests
import tempfile
import json
from urllib.parse import urlparse
import time

# Import text extraction libraries
try:
    import PyPDF2
    import pdfplumber
    from docx import Document
except ImportError as e:
    print(f"Missing required library: {e}")
    sys.exit(1)

class CVProcessor:
    def __init__(self, file_url, application_id, callback_url, auth_token):
        self.file_url = file_url
        self.application_id = application_id
        self.callback_url = callback_url
        self.auth_token = auth_token
        self.temp_file_path = None
        
    def download_file(self):
        """Download CV file with proper authentication headers"""
        print(f"Downloading file from: {self.file_url}")
        
        headers = {
            'Authorization': f'Bearer {self.auth_token}',
            'User-Agent': 'GitHub-Actions-CV-Processor/1.0'
        }
        
        try:
            # Create a temporary file
            temp_fd, self.temp_file_path = tempfile.mkstemp()
            os.close(temp_fd)  # Close the file descriptor, keep the path
            
            # Download with authentication
            response = requests.get(
                self.file_url, 
                headers=headers, 
                stream=True, 
                timeout=60,
                allow_redirects=True
            )
            
            response.raise_for_status()
            
            # Determine file extension from Content-Type or URL
            content_type = response.headers.get('content-type', '').lower()
            file_extension = self._get_file_extension(content_type)
            
            # Add proper extension to temp file
            if file_extension:
                new_temp_path = self.temp_file_path + file_extension
                os.rename(self.temp_file_path, new_temp_path)
                self.temp_file_path = new_temp_path
            
            # Write file content
            with open(self.temp_file_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            print(f"File downloaded successfully to: {self.temp_file_path}")
            return True
            
        except requests.exceptions.RequestException as e:
            print(f"Error downloading file: {e}")
            return False
        except Exception as e:
            print(f"Unexpected error during download: {e}")
            return False
    
    def _get_file_extension(self, content_type):
        """Get file extension based on content type"""
        content_type_map = {
            'application/pdf': '.pdf',
            'application/msword': '.doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx'
        }
        
        for mime_type, extension in content_type_map.items():
            if mime_type in content_type:
                return extension
        
        # Fallback to URL extension
        parsed_url = urlparse(self.file_url)
        path = parsed_url.path
        if '.' in path:
            return os.path.splitext(path)[1]
        
        return None
    
    def extract_text(self):
        """Extract text from the downloaded file"""
        if not self.temp_file_path or not os.path.exists(self.temp_file_path):
            raise Exception("No file available for text extraction")
        
        file_extension = os.path.splitext(self.temp_file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_pdf_text()
            elif file_extension == '.docx':
                return self._extract_docx_text()
            elif file_extension == '.doc':
                return self._extract_doc_text()
            else:
                raise Exception(f"Unsupported file type: {file_extension}")
        except Exception as e:
            print(f"Error extracting text: {e}")
            raise
    
    def _extract_pdf_text(self):
        """Extract text from PDF using multiple methods"""
        text = ""
        
        # Try pdfplumber first (better for complex PDFs)
        try:
            with pdfplumber.open(self.temp_file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                print("PDF text extracted using pdfplumber")
                return text
        except Exception as e:
            print(f"pdfplumber failed: {e}")
        
        # Fallback to PyPDF2
        try:
            with open(self.temp_file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            if text.strip():
                print("PDF text extracted using PyPDF2")
                return text
        except Exception as e:
            print(f"PyPDF2 failed: {e}")
        
        # Last resort: try pdftotext command
        try:
            import subprocess
            result = subprocess.run(
                ['pdftotext', self.temp_file_path, '-'], 
                capture_output=True, 
                text=True, 
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                print("PDF text extracted using pdftotext")
                return result.stdout
        except Exception as e:
            print(f"pdftotext failed: {e}")
        
        raise Exception("All PDF text extraction methods failed")
    
    def _extract_docx_text(self):
        """Extract text from DOCX file"""
        document = Document(self.temp_file_path)
        text = ""
        
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables if any
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        print("DOCX text extracted successfully")
        return text
    
    def _extract_doc_text(self):
        """Extract text from DOC file using antiword"""
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', self.temp_file_path], 
                capture_output=True, 
                text=True, 
                timeout=30
            )
            if result.returncode == 0:
                print("DOC text extracted using antiword")
                return result.stdout
            else:
                raise Exception(f"antiword failed with return code {result.returncode}")
        except FileNotFoundError:
            raise Exception("antiword command not found. Please install antiword package.")
        except Exception as e:
            raise Exception(f"Failed to extract DOC text: {e}")
    
    def send_callback(self, extracted_text, success=True, error_message=None):
        """Send processing results back to Laravel application"""
        print(f"Sending callback to: {self.callback_url}")
        
        # Get current timestamp
        try:
            response = requests.get('https://worldtimeapi.org/api/timezone/Etc/UTC', timeout=10)
            timestamp = response.json().get('datetime', time.strftime('%Y-%m-%dT%H:%M:%SZ'))
        except:
            timestamp = time.strftime('%Y-%m-%dT%H:%M:%SZ')
        
        payload = {
            'application_id': self.application_id,
            'success': success,
            'timestamp': timestamp
        }
        
        if success:
            payload['extracted_text'] = extracted_text
            payload['text_length'] = len(extracted_text)
        else:
            payload['error'] = error_message or "Unknown error occurred"
        
        headers = {
            'Authorization': f'Bearer {self.auth_token}',
            'Content-Type': 'application/json',
            'User-Agent': 'GitHub-Actions-CV-Processor/1.0'
        }
        
        try:
            response = requests.post(
                self.callback_url,
                json=payload,
                headers=headers,
                timeout=30
            )
            
            response.raise_for_status()
            print(f"Callback sent successfully. Status: {response.status_code}")
            return True
            
        except Exception as e:
            print(f"Failed to send callback: {e}")
            return False
    
    def cleanup(self):
        """Clean up temporary files"""
        if self.temp_file_path and os.path.exists(self.temp_file_path):
            try:
                os.unlink(self.temp_file_path)
                print(f"Cleaned up temporary file: {self.temp_file_path}")
            except Exception as e:
                print(f"Failed to clean up temporary file: {e}")
    
    def process(self):
        """Main processing method"""
        try:
            # Download the file
            if not self.download_file():
                self.send_callback(None, False, "Failed to download CV file")
                return False
            
            # Extract text
            extracted_text = self.extract_text()
            
            if not extracted_text.strip():
                self.send_callback(None, False, "No text could be extracted from the CV file")
                return False
            
            # Log success (don't log full content for privacy)
            print(f"Text extraction successful. Length: {len(extracted_text)} characters")
            print(f"Preview: {extracted_text[:100]}...")
            
            # Send success callback
            return self.send_callback(extracted_text, True)
            
        except Exception as e:
            print(f"Processing error: {e}")
            self.send_callback(None, False, str(e))
            return False
        
        finally:
            self.cleanup()

def main():
    """Main entry point"""
    if len(sys.argv) != 5:
        print("Usage: process_cv.py <file_url> <application_id> <callback_url> <auth_token>")
        sys.exit(1)
    
    file_url = sys.argv[1]
    application_id = sys.argv[2]
    callback_url = sys.argv[3]
    auth_token = sys.argv[4]
    
    print(f"Starting CV processing for application ID: {application_id}")
    
    processor = CVProcessor(file_url, application_id, callback_url, auth_token)
    success = processor.process()
    
    if success:
        print("CV processing completed successfully")
        sys.exit(0)
    else:
        print("CV processing failed")
        sys.exit(1)

if __name__ == "__main__":
    main()